#!/usr/bin/env python
"""Automate bundling of Word and PDF documents into a single PDF.

This tool scans the current directory for ``.docx`` and ``.pdf`` files,
creates cover pages, builds a table of contents and merges everything
into ``final_output.pdf``. Microsoft Word is required for converting
``.docx`` files to PDF, so the script is intended for Windows systems.
"""

import sys
import time
import logging
from pathlib import Path
from typing import List, Tuple, Dict

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
import PyPDF2
import fitz  # PyMuPDF

###############################################################################
# Configuration & Constants
###############################################################################
SCRIPT_DIR = Path(__file__).parent.absolute()
OUTPUT_DIR = SCRIPT_DIR / "output"
LOG_FILE = SCRIPT_DIR / "script_log.txt"
FINAL_PDF = SCRIPT_DIR / "final_output.pdf"

BATES_START = 1         # Final PDF page numbering will start at 001
BATES_FONT_SIZE = 14
TIMEOUT_BASE = 1.0

###############################################################################
# Logging Setup
###############################################################################
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [LOG] %(levelname)s: %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8", mode='w'),
        logging.StreamHandler(sys.stdout)
    ]
)
log = logging.info
log_error = logging.error
log_debug = logging.debug

if sys.platform != "win32":
    log_error(
        "This script runs only on Windows because it uses Microsoft Word for"
        " DOCX to PDF conversion."
    )
    sys.exit(1)

try:
    from comtypes.client import CreateObject
except Exception as exc:  # pragma: no cover - platform specific
    log_error(
        "Microsoft Word COM automation is unavailable. Ensure Word is\n"
        "installed correctly: %s",
        exc,
    )
    sys.exit(1)

###############################################################################
# Dependency Checks
###############################################################################
REQUIRED_PACKAGES = ["docx", "PyPDF2", "fitz"]


def verify_dependencies() -> None:
    """Ensure all required packages are available."""
    missing = []
    for pkg in REQUIRED_PACKAGES:
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pkg)
    if missing:
        log_error(
            "Missing required packages: %s. "
            "Install them with 'pip install -r requirements.txt'."
            % ", ".join(missing)
        )
        sys.exit(1)


verify_dependencies()

###############################################################################
# Helper Functions
###############################################################################


def remove_table_borders(table: docx.table.Table) -> None:
    """Remove table-level borders from a docx table."""
    for border in table._element.xpath('.//w:tblBorders'):
        border.getparent().remove(border)


def remove_cell_borders(table: docx.table.Table) -> None:
    """Remove cell-level borders from a docx table."""
    for row in table.rows:
        for cell in row.cells:
            for border in cell._tc.xpath('.//w:tcBorders'):
                border.getparent().remove(border)


def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    """
    Set margins (in dxa units) for a table cell.
    1 dxa = 1/20th of a point.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for margin, value in (
        ('top', top),
        ('left', left),
        ('bottom', bottom),
        ('right', right),
    ):
        element = tcMar.find(qn(f'w:{margin}'))
        if element is None:
            element = OxmlElement(f'w:{margin}')
            tcMar.append(element)
        element.set(qn('w:w'), str(value))
        element.set(qn('w:type'), 'dxa')


def apply_document_styles(doc: docx.Document) -> None:
    """Apply consistent styling (font, margins, spacing) to the document."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert a DOCX file to PDF using Microsoft Word COM automation."""
    log(f"Converting {docx_path} -> {pdf_path}")
    if not docx_path.exists():
        raise FileNotFoundError(f"{docx_path} not found.")
    word = None
    try:
        word = CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(str(docx_path), ReadOnly=True)
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()
        time.sleep(TIMEOUT_BASE)
    except Exception as e:
        log_error(f"Error converting {docx_path}: {e}")
        raise
    finally:
        if word is not None:
            word.Quit()
    if not pdf_path.exists():
        raise FileNotFoundError(f"Conversion failed: {pdf_path} not created.")
    log(f"Converted {docx_path.name} -> {pdf_path.name}")


def pdf_page_count(pdf_path: Path) -> int:
    """Return the number of pages in a PDF file."""
    try:
        reader = PyPDF2.PdfReader(str(pdf_path))
        return len(reader.pages)
    except Exception as e:
        log_error(f"Error counting pages in {pdf_path}: {e}")
        return 0


def merge_pdfs(pdf_paths: List[Path], final_pdf: Path) -> None:
    """Merge PDFs in the provided order into a final PDF."""
    from PyPDF2 import PdfMerger
    log(f"Merging PDFs into {final_pdf.name}")
    merger = PdfMerger()
    for p in pdf_paths:
        if p.exists():
            pages = pdf_page_count(p)
            if pages > 0:
                log(f"Appending {p.name} ({pages} pages)")
                merger.append(str(p))
            else:
                log(f"Skipping {p.name}, 0 pages.")
        else:
            log_error(f"File missing: {p.name}")
    merger.write(str(final_pdf))
    merger.close()
    if not final_pdf.exists():
        raise FileNotFoundError(f"Merged PDF not created: {final_pdf}")
    log(f"Merged PDF saved: {final_pdf.name}")


def apply_bates_numbering(
    pdf_path: Path,
    start_number: int = BATES_START,
    font_size: int = BATES_FONT_SIZE,
) -> None:
    """Apply sequential Bates numbering with a subtle background
    to each page."""
    log(f"Applying Bates numbering to {pdf_path.name}")
    doc = fitz.open(str(pdf_path))
    for i, page in enumerate(doc):
        bates_num = f"{start_number + i:03d}"
        rect = page.rect
        x = rect.width - 80
        y = rect.height - 40
        page.draw_rect(
            fitz.Rect(x - 10, y - 25, x + 60, y + 5),
            color=(0.9, 0.9, 0.9),
            fill=(0.9, 0.9, 0.9)
        )
        page.insert_text((x, y), bates_num,
                         fontname="Helvetica",
                         fontsize=font_size,
                         color=(0, 0, 0))
    temp_pdf = pdf_path.with_suffix('.temp.pdf')
    doc.save(str(temp_pdf))
    doc.close()
    temp_pdf.replace(pdf_path)
    log(f"Bates numbering applied to {pdf_path.name}")


def add_pdf_bookmarks(pdf_path: Path, toc_list: List[List]) -> None:
    """
    Add bookmarks (outline) to the merged PDF.
    Each entry in toc_list is [level, title, page_number] (1-indexed).
    """
    log("Adding PDF bookmarks (outline)...")
    doc = fitz.open(str(pdf_path))
    doc.set_toc(toc_list)
    new_toc = doc.get_toc()
    log_debug(f"New TOC: {new_toc}")
    doc.save(str(pdf_path))
    doc.close()
    log("PDF bookmarks added.")

###############################################################################
# Document Creation Functions
###############################################################################


def create_cover_page(cover_docx: Path, number: str, file_name: str) -> None:
    """
    Generate a cover page with a decorative header,
    document number, and filename.
    The date is removed from cover pages.
    """
    log(f"Creating cover page: {cover_docx.name}")
    doc = docx.Document()
    apply_document_styles(doc)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("DOCUMENT INDEX")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    for _ in range(6):
        doc.add_paragraph()
    num_para = doc.add_paragraph()
    num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_run = num_para.add_run(f" {number} ")
    num_run.font.name = "Arial"
    num_run.font.size = Pt(18)
    num_run.font.bold = True
    num_para.paragraph_format.space_after = Pt(12)
    file_para = doc.add_paragraph()
    file_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_run = file_para.add_run(file_name)
    file_run.font.name = "Arial"
    file_run.font.size = Pt(24)
    file_run.font.bold = True
    file_para.paragraph_format.space_after = Pt(24)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("⎯" * 20)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)
    doc.save(str(cover_docx))


def create_contents_page(
    contents_docx: Path,
    all_items: List[Tuple[str, Path]],
    bates_map: Dict[str, int],
) -> None:
    """
    Create a formatted table of contents with document names, Bates numbers,
    and the generation date in the header.
    """
    log(f"Creating contents page: {contents_docx.name}")
    doc = docx.Document()
    apply_document_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TABLE OF CONTENTS")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(20)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(time.strftime("%B %d, %Y"))
    date_run.font.name = "Arial"
    date_run.font.size = Pt(12)
    date_run.font.italic = True
    date_para.paragraph_format.space_after = Pt(20)
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider.add_run("⎯" * 30)
    divider_run.font.color.rgb = RGBColor(180, 180, 180)
    divider.paragraph_format.space_after = Pt(20)
    table = doc.add_table(rows=len(all_items) + 1, cols=2)
    table.style = "Table Grid"
    remove_table_borders(table)
    remove_cell_borders(table)
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    widths = [Inches(6.0), Inches(1.0)]
    for col_idx, width in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.cells[0].text = "Document"
    header_row.cells[1].text = "Page"
    for cell in header_row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=100, bottom=100)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
    header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, (num, path_obj) in enumerate(all_items):
        row_idx = i + 1
        indent = "    " * num.count(".")
        left_text = f"{indent}{num} - {path_obj.name}"
        cell_left = table.cell(row_idx, 0)
        cell_left.text = ""
        cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell_left, top=100, bottom=100)
        p_left = cell_left.paragraphs[0]
        run_left = p_left.add_run(left_text)
        run_left.font.name = "Arial"
        run_left.font.size = Pt(11)
        if path_obj.is_dir():
            run_left.font.bold = True
        cell_right = table.cell(row_idx, 1)
        cell_right.text = ""
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell_right, top=100, bottom=100)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if num in bates_map and path_obj.is_file():
            run_right = p_right.add_run(f"{bates_map[num]:03d}")
            run_right.font.name = "Arial"
            run_right.font.size = Pt(11)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("")
    doc.save(str(contents_docx))
    log(f"Contents page saved: {contents_docx.name}")

###############################################################################
# Main Execution Flow
###############################################################################


def main() -> None:
    try:
        OUTPUT_DIR.mkdir(exist_ok=True)
        log(f"Output directory ready: {OUTPUT_DIR}")
        all_items: List[Tuple[str, Path]] = []

        def scan_dir(dirpath: Path, prefix: str = ""):
            if OUTPUT_DIR in dirpath.parents or dirpath == OUTPUT_DIR:
                return
            files = sorted(
                [
                    f
                    for f in dirpath.iterdir()
                    if (
                        f.is_file()
                        and f.suffix.lower() in [".docx", ".pdf"]
                        and f != FINAL_PDF
                    )
                ],
                key=lambda x: x.name.lower(),
            )
            i = 1
            for f in files:
                num = f"{prefix}{i}"
                all_items.append((num, f))
                i += 1
            subdirs = sorted(
                [
                    d
                    for d in dirpath.iterdir()
                    if d.is_dir() and d != OUTPUT_DIR
                ],
                key=lambda x: x.name.lower(),
            )
            j = 1
            for d in subdirs:
                sub_prefix = f"{prefix}{j}."
                all_items.append((sub_prefix[:-1], d))
                j += 1
                scan_dir(d, prefix=sub_prefix)
        scan_dir(SCRIPT_DIR, "")
        real_files: List[Tuple[str, Path]] = [
            (num, p) for num, p in all_items
            if (
                p.is_file()
                and p.suffix.lower() in [".docx", ".pdf"]
                and p != FINAL_PDF
            )
        ]
        processed: Dict[str, Tuple[Path, int, Path, int]] = {}
        for num, p in real_files:
            cover_docx = OUTPUT_DIR / f"cover_{num}.docx"
            cover_pdf = OUTPUT_DIR / f"cover_{num}.pdf"
            create_cover_page(cover_docx, num, p.name)
            convert_docx_to_pdf(cover_docx, cover_pdf)
            cover_pages = pdf_page_count(cover_pdf)
            if p.suffix.lower() == ".docx":
                file_pdf = OUTPUT_DIR / f"file_{num}.pdf"
                convert_docx_to_pdf(p, file_pdf)
            else:
                file_pdf = p
            file_pages = pdf_page_count(file_pdf)
            processed[num] = (cover_pdf, cover_pages, file_pdf, file_pages)
            log_debug(
                f"Processed {num}: cover={cover_pdf.name} "
                f"({cover_pages} pages), "
                f"file={file_pdf.name} ({file_pages} pages)"
            )
        dummy_contents_docx = OUTPUT_DIR / "contents_dummy.docx"
        create_contents_page(dummy_contents_docx, all_items, {})
        dummy_contents_pdf = OUTPUT_DIR / "contents_dummy.pdf"
        convert_docx_to_pdf(dummy_contents_docx, dummy_contents_pdf)
        contents_pages = pdf_page_count(dummy_contents_pdf)
        log_debug(f"Dummy contents PDF pages: {contents_pages}")
        current_page = BATES_START + contents_pages
        bates_map: Dict[str, int] = {}
        for num, _ in real_files:
            cover_pdf, cover_pages, file_pdf, file_pages = processed[num]
            bates_map[num] = current_page
            current_page += cover_pages + file_pages
        log_debug(f"Bates mapping: {bates_map}")
        contents_docx = OUTPUT_DIR / "contents.docx"
        create_contents_page(contents_docx, all_items, bates_map)
        contents_pdf = OUTPUT_DIR / "contents.pdf"
        convert_docx_to_pdf(contents_docx, contents_pdf)
        final_order: List[Path] = [contents_pdf]
        for num, _ in real_files:
            cover_pdf, _, file_pdf, _ = processed[num]
            final_order.extend([cover_pdf, file_pdf])
        merge_pdfs(final_order, FINAL_PDF)
        apply_bates_numbering(
            FINAL_PDF,
            start_number=BATES_START,
            font_size=BATES_FONT_SIZE,
        )
        toc_list = []
        for num, p in real_files:
            title = f"{num} - {p.name}"
            toc_list.append([1, title, bates_map[num]])
        add_pdf_bookmarks(FINAL_PDF, toc_list)
        for pattern in [
            "cover_*.docx",
            "cover_*.pdf",
            "file_*.pdf",
            "contents_dummy.*",
            "contents.docx",
            "contents.pdf",
        ]:
            for temp in OUTPUT_DIR.glob(pattern):
                try:
                    temp.unlink()
                    log_debug(f"Removed temporary file: {temp}")
                except Exception:
                    pass
        log("Process complete. Check final_output.pdf and script_log.txt.")
    except Exception as e:
        log_error(f"Fatal error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
