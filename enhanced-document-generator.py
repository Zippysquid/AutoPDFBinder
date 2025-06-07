#!/usr/bin/env python
from __future__ import annotations

import os
import sys
import subprocess
import time
import logging
import shutil
from pathlib import Path
from typing import List, Tuple, Dict

###############################################################################
# Configuration & Constants
###############################################################################
SCRIPT_DIR = Path(__file__).parent.absolute()
OUTPUT_DIR = SCRIPT_DIR / "output"
LOG_FILE = SCRIPT_DIR / "script_log.txt"
FINAL_PDF = SCRIPT_DIR / "final_output.pdf"

BATES_START = 1         # Final PDF page numbering will start at 001
BATES_FONT_SIZE = 14
TIMEOUT_BASE = 1.0

###############################################################################
# Logging Setup
###############################################################################
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [LOG] %(levelname)s: %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8", mode='w'),
        logging.StreamHandler(sys.stdout)
    ]
)
log = logging.info
log_error = logging.error
log_debug = logging.debug

###############################################################################
# Dependency Installation
###############################################################################
def install_missing_packages() -> None:
    required_packages = {
        "python-docx": "docx",
        "pypdf2": "PyPDF2",
        "pymupdf": "fitz",
        "comtypes": "comtypes"
    }
    log(f"Checking dependencies with Python: {sys.executable}")
    missing = False
    for pkg_name, import_name in required_packages.items():
        try:
            __import__(import_name)
            log(f"Package {pkg_name} ({import_name}) is installed.")
        except ImportError:
            log(f"Installing {pkg_name}...")
            try:
                subprocess.run(
                    [sys.executable, "-m", "pip", "install", "--user", pkg_name],
                    capture_output=True, text=True, check=True, timeout=300
                )
                log(f"Installed {pkg_name}")
            except subprocess.CalledProcessError as e:
                log_error(f"Failed to install {pkg_name}: {e.stderr}")
                missing = True
    if missing:
        log_error("Critical: Some packages failed to install.")
        sys.exit(1)


###############################################################################
# Helper Functions
###############################################################################
def remove_table_borders(table: docx.table.Table) -> None:
    """Remove table-level borders from a docx table."""
    for border in table._element.xpath('.//w:tblBorders'):
        border.getparent().remove(border)

def remove_cell_borders(table: docx.table.Table) -> None:
    """Remove cell-level borders from a docx table."""
    for row in table.rows:
        for cell in row.cells:
            for border in cell._tc.xpath('.//w:tcBorders'):
                border.getparent().remove(border)

def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    """
    Set margins (in dxa units) for a table cell.
    1 dxa = 1/20th of a point.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for margin, value in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        element = tcMar.find(qn(f'w:{margin}'))
        if element is None:
            element = OxmlElement(f'w:{margin}')
            tcMar.append(element)
        element.set(qn('w:w'), str(value))
        element.set(qn('w:type'), 'dxa')

def apply_document_styles(doc: "docx.Document") -> None:
    """Apply consistent styling (font, margins, spacing) to the document."""
    from docx.shared import Pt, Inches

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_centered_paragraph(doc: "docx.Document", text: str, size: int,
                           bold: bool = False, italic: bool = False,
                           space_after: int | None = None,
                           color: tuple[int, int, int] | None = None):
    """Helper to add a centered paragraph with consistent styling."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert a DOCX file to PDF using Microsoft Word COM automation."""
    from comtypes.client import CreateObject

    log(f"Converting {docx_path} -> {pdf_path}")
    if not docx_path.exists():
        raise FileNotFoundError(f"{docx_path} not found.")
    word = None
    try:
        word = CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(str(docx_path), ReadOnly=True)
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()
        time.sleep(TIMEOUT_BASE)
    except Exception as e:
        log_error(f"Error converting {docx_path}: {e}")
        raise
    finally:
        if word is not None:
            word.Quit()
    if not pdf_path.exists():
        raise FileNotFoundError(f"Conversion failed: {pdf_path} not created.")
    log(f"Converted {docx_path.name} -> {pdf_path.name}")

def pdf_page_count(pdf_path: Path) -> int:
    """Return the number of pages in a PDF file."""
    import PyPDF2
    try:
        reader = PyPDF2.PdfReader(str(pdf_path))
        return len(reader.pages)
    except Exception as e:
        log_error(f"Error counting pages in {pdf_path}: {e}")
        return 0

def merge_pdfs(pdf_paths: List[Path], final_pdf: Path) -> None:
    """Merge PDFs in the provided order into a final PDF."""
    from PyPDF2 import PdfMerger
    log(f"Merging PDFs into {final_pdf.name}")
    merger = PdfMerger()
    for p in pdf_paths:
        if p.exists():
            pages = pdf_page_count(p)
            if pages > 0:
                log(f"Appending {p.name} ({pages} pages)")
                merger.append(str(p))
            else:
                log(f"Skipping {p.name}, 0 pages.")
        else:
            log_error(f"File missing: {p.name}")
    merger.write(str(final_pdf))
    merger.close()
    if not final_pdf.exists():
        raise FileNotFoundError(f"Merged PDF not created: {final_pdf}")
    log(f"Merged PDF saved: {final_pdf.name}")

def apply_bates_numbering(pdf_path: Path, start_number: int = BATES_START, font_size: int = BATES_FONT_SIZE) -> None:
    """Apply sequential Bates numbering with a subtle background to each page."""
    import fitz

    log(f"Applying Bates numbering to {pdf_path.name}")
    doc = fitz.open(str(pdf_path))
    for i, page in enumerate(doc):
        bates_num = f"{start_number + i:03d}"
        rect = page.rect
        x = rect.width - 80
        y = rect.height - 40
        page.draw_rect(
            fitz.Rect(x - 10, y - 25, x + 60, y + 5),
            color=(0.9, 0.9, 0.9),
            fill=(0.9, 0.9, 0.9)
        )
        page.insert_text((x, y), bates_num,
                         fontname="Helvetica",
                         fontsize=font_size,
                         color=(0, 0, 0))
    temp_pdf = pdf_path.with_suffix('.temp.pdf')
    doc.save(str(temp_pdf))
    doc.close()
    temp_pdf.replace(pdf_path)
    log(f"Bates numbering applied to {pdf_path.name}")

def add_pdf_bookmarks(pdf_path: Path, toc_list: List[List]) -> None:
    """
    Add bookmarks (outline) to the merged PDF.
    Each entry in toc_list is [level, title, page_number] (1-indexed).
    """
    import fitz

    log("Adding PDF bookmarks (outline)...")
    doc = fitz.open(str(pdf_path))
    doc.set_toc(toc_list)
    new_toc = doc.get_toc()
    log_debug(f"New TOC: {new_toc}")
    temp_pdf = pdf_path.with_suffix('.bm.pdf')
    doc.save(str(temp_pdf))
    doc.close()
    temp_pdf.replace(pdf_path)
    log("PDF bookmarks added.")

def add_toc_links(pdf_path: Path, link_entries: List[Tuple[str, int]], contents_pages: int) -> None:
    """Insert clickable links on the table of contents pages."""
    import fitz

    log("Adding clickable links to contents page(s)...")
    doc = fitz.open(str(pdf_path))
    for text, target in link_entries:
        found = False
        for i in range(contents_pages):
            page = doc[i]
            rects = page.search_for(text)
            for r in rects:
                link = {"kind": fitz.LINK_GOTO, "page": target - 1, "from": r}
                page.insert_link(link)
                found = True
        if not found:
            log_debug(f"Link text not found: {text}")
    temp_pdf = pdf_path.with_suffix('.links.pdf')
    doc.save(str(temp_pdf))
    doc.close()
    temp_pdf.replace(pdf_path)
    log("Contents links added.")

###############################################################################
# Document Creation Functions
###############################################################################
def create_cover_page(cover_docx: Path, number: str, file_name: str) -> None:
    """
    Generate a cover page with a decorative header, document number, and filename.
    The date is removed from cover pages.
    """
    import docx

    log(f"Creating cover page: {cover_docx.name}")
    doc = docx.Document()
    apply_document_styles(doc)

    add_centered_paragraph(doc, "DOCUMENT INDEX", 12, bold=True,
                           color=(100, 100, 100))
    for _ in range(6):
        doc.add_paragraph()

    add_centered_paragraph(doc, f" {number} ", 18, bold=True,
                           space_after=12)
    add_centered_paragraph(doc, file_name, 24, bold=True,
                           space_after=24)
    add_centered_paragraph(doc, "⎯" * 20, 11, color=(100, 100, 100))

    doc.save(str(cover_docx))

def create_contents_page(contents_docx: Path, all_items: List[Tuple[str, Path]], bates_map: Dict[str, int]) -> None:
    """
    Create a formatted table of contents with document names, Bates numbers,
    and the generation date in the header.
    """
    import docx
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor

    log(f"Creating contents page: {contents_docx.name}")
    doc = docx.Document()
    apply_document_styles(doc)

    add_centered_paragraph(doc, "TABLE OF CONTENTS", 16, bold=True,
                           space_after=20)
    add_centered_paragraph(doc, time.strftime("%B %d, %Y"), 12,
                           italic=True, space_after=20)
    add_centered_paragraph(doc, "⎯" * 30, 11, color=(180, 180, 180),
                           space_after=20)
    table = doc.add_table(rows=len(all_items) + 1, cols=2)
    table.style = "Table Grid"
    remove_table_borders(table)
    remove_cell_borders(table)
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    widths = [Inches(5.3), Inches(1.2)]
    for col_idx, width in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.cells[0].text = "Document"
    header_row.cells[1].text = "Page"
    for cell in header_row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_margin = 200 if cell is header_row.cells[1] else 0
        set_cell_margins(cell, top=100, bottom=100, right=right_margin)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
    header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, (num, path_obj) in enumerate(all_items):
        row_idx = i + 1
        indent = "    " * num.count(".")
        left_text = f"{indent}{num} - {path_obj.name}"
        cell_left = table.cell(row_idx, 0)
        cell_left.text = ""
        cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell_left, top=100, bottom=100)
        p_left = cell_left.paragraphs[0]
        run_left = p_left.add_run(left_text)
        run_left.font.name = "Arial"
        run_left.font.size = Pt(11)
        if path_obj.is_dir():
            run_left.font.bold = True
        cell_right = table.cell(row_idx, 1)
        cell_right.text = ""
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell_right, top=100, bottom=100, right=200)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if num in bates_map and path_obj.is_file():
            run_right = p_right.add_run(f"{bates_map[num]:03d}")
            run_right.font.name = "Arial"
            run_right.font.size = Pt(11)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("")
    doc.save(str(contents_docx))
    log(f"Contents page saved: {contents_docx.name}")

###############################################################################
# Main Execution Flow
###############################################################################
def scan_dir(dirpath: Path, prefix: str = "", collected: List[Tuple[str, Path]] | None = None) -> List[Tuple[str, Path]]:
    """Recursively scan a directory for docx and pdf files."""
    if collected is None:
        collected = []
    if OUTPUT_DIR in dirpath.parents or dirpath == OUTPUT_DIR:
        return collected
    files = sorted(
        [f for f in dirpath.iterdir() if f.is_file() and f.suffix.lower() in [".docx", ".pdf"] and f != FINAL_PDF],
        key=lambda x: x.name.lower(),
    )
    i = 1
    for f in files:
        num = f"{prefix}{i}"
        collected.append((num, f))
        i += 1
    subdirs = sorted(
        [d for d in dirpath.iterdir() if d.is_dir() and d != OUTPUT_DIR],
        key=lambda x: x.name.lower(),
    )
    j = 1
    for d in subdirs:
        sub_prefix = f"{prefix}{j}."
        collected.append((sub_prefix[:-1], d))
        j += 1
        scan_dir(d, prefix=sub_prefix, collected=collected)
    return collected


def process_files(real_files: List[Tuple[str, Path]]) -> Dict[str, Tuple[Path, int, Path, int]]:
    """Create cover pages, convert documents and gather page counts."""
    processed: Dict[str, Tuple[Path, int, Path, int]] = {}
    for num, p in real_files:
        cover_docx = OUTPUT_DIR / f"cover_{num}.docx"
        cover_pdf = OUTPUT_DIR / f"cover_{num}.pdf"
        create_cover_page(cover_docx, num, p.name)
        convert_docx_to_pdf(cover_docx, cover_pdf)
        cover_pages = pdf_page_count(cover_pdf)
        if p.suffix.lower() == ".docx":
            file_pdf = OUTPUT_DIR / f"file_{num}.pdf"
            convert_docx_to_pdf(p, file_pdf)
        else:
            file_pdf = p
        file_pages = pdf_page_count(file_pdf)
        processed[num] = (cover_pdf, cover_pages, file_pdf, file_pages)
        log_debug(
            f"Processed {num}: cover={cover_pdf.name} ({cover_pages} pages), file={file_pdf.name} ({file_pages} pages)"
        )
    return processed


def build_bates_map(real_files: List[Tuple[str, Path]], contents_pages: int, processed: Dict[str, Tuple[Path, int, Path, int]]) -> Dict[str, int]:
    """Calculate Bates start page for each document."""
    current_page = BATES_START + contents_pages
    bates_map: Dict[str, int] = {}
    for num, _ in real_files:
        cover_pdf, cover_pages, file_pdf, file_pages = processed[num]
        bates_map[num] = current_page
        current_page += cover_pages + file_pages
    log_debug(f"Bates mapping: {bates_map}")
    return bates_map


def cleanup_temp_files() -> None:
    """Remove intermediate files in the output directory."""
    for pattern in ["cover_*.docx", "cover_*.pdf", "file_*.pdf", "contents_dummy.*", "contents.docx", "contents.pdf"]:
        for temp in OUTPUT_DIR.glob(pattern):
            try:
                temp.unlink()
                log_debug(f"Removed temporary file: {temp}")
            except Exception:
                pass


def main() -> None:
    try:
        OUTPUT_DIR.mkdir(exist_ok=True)
        log(f"Output directory ready: {OUTPUT_DIR}")
        install_missing_packages()

        all_items = scan_dir(SCRIPT_DIR)
        real_files: List[Tuple[str, Path]] = [
            (num, p) for num, p in all_items
            if p.is_file() and p.suffix.lower() in [".docx", ".pdf"] and p != FINAL_PDF
        ]

        processed = process_files(real_files)

        dummy_contents_docx = OUTPUT_DIR / "contents_dummy.docx"
        create_contents_page(dummy_contents_docx, all_items, {})
        dummy_contents_pdf = OUTPUT_DIR / "contents_dummy.pdf"
        convert_docx_to_pdf(dummy_contents_docx, dummy_contents_pdf)
        contents_pages = pdf_page_count(dummy_contents_pdf)
        log_debug(f"Dummy contents PDF pages: {contents_pages}")

        bates_map = build_bates_map(real_files, contents_pages, processed)
        contents_docx = OUTPUT_DIR / "contents.docx"
        create_contents_page(contents_docx, all_items, bates_map)
        contents_pdf = OUTPUT_DIR / "contents.pdf"
        convert_docx_to_pdf(contents_docx, contents_pdf)
        final_order: List[Path] = [contents_pdf]
        for num, _ in real_files:
            cover_pdf, _, file_pdf, _ = processed[num]
            final_order.extend([cover_pdf, file_pdf])
        merge_pdfs(final_order, FINAL_PDF)
        apply_bates_numbering(FINAL_PDF, start_number=BATES_START, font_size=BATES_FONT_SIZE)
        toc_list = []
        link_entries = []
        for num, p in real_files:
            title = f"{num} - {p.name}"
            toc_list.append([1, title, bates_map[num]])
            indent = "    " * num.count(".")
            link_entries.append((f"{indent}{title}", bates_map[num]))
        add_pdf_bookmarks(FINAL_PDF, toc_list)
        add_toc_links(FINAL_PDF, link_entries, contents_pages)
        cleanup_temp_files()
        log("Process complete. Check final_output.pdf and script_log.txt.")
    except Exception as e:
        log_error(f"Fatal error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
